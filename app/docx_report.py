# app/docx_report.py
# ==============================================================================
# Tujuan : Generate laporan sensor status dan insiden dalam format DOCX
# Caller : app/internal_api.py (via generate_report endpoint)
# ==============================================================================
import io
from datetime import datetime
from typing import List, Dict, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_kv_table(doc: Document, rows: List[List[str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value


def _add_data_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val)


def _gas_status(value: float, threshold: float) -> str:
    if value == 0:
        return "Offline / Belum Terhubung"
    return f"TINGGI (>={threshold})" if value >= threshold else "Normal"


def _temp_status(temp: float) -> str:
    if temp == 0:
        return "Offline / Belum Terhubung"
    if temp >= 60:
        return "SANGAT TINGGI"
    if temp >= 45:
        return "TINGGI"
    return "Normal"


def _sensor_recs(cameras_data: Dict) -> List[str]:
    all_data = list(cameras_data.values())
    recs = []
    if any(d.get("mq2", 1) == 0 or d.get("mq7", 1) == 0 for d in all_data):
        recs.append("Periksa koneksi ESP32 — beberapa sensor menunjukkan nilai 0 (offline).")
    recs += [
        "Lakukan kalibrasi sensor MQ secara berkala sesuai jadwal perawatan.",
        "Pastikan APAR di sekitar area sensor dalam kondisi siap pakai.",
        "Verifikasi jalur evakuasi bebas hambatan.",
        "Pertahankan ventilasi ruangan agar pembacaan sensor tetap akurat.",
    ]
    return recs


def _incident_recs(alerts: List[Dict], anomaly_count: int) -> List[str]:
    recs = []
    if alerts:
        recs += [
            "Lakukan investigasi fisik pada area kamera yang mendeteksi alert.",
            "Dokumentasikan penyebab insiden dan tindakan yang diambil.",
            "Pastikan sistem notifikasi berfungsi dan tim respons dapat dihubungi.",
        ]
    if anomaly_count > 0:
        recs += [
            "Review data sensor pada waktu anomali untuk identifikasi pola.",
            "Pertimbangkan penyesuaian threshold LSTM jika anomali bersifat false positive.",
        ]
    recs += [
        "Lakukan simulasi evakuasi rutin sesuai prosedur K3.",
        "Pastikan baterai backup sistem alarm berfungsi dengan baik.",
    ]
    return recs


def generate_sensor_report(sensor_data: Dict, camera_id: str = "all") -> bytes:
    """Return DOCX bytes for a sensor status report."""
    doc = Document()
    now_str = datetime.now().strftime("%d %B %Y, %H:%M:%S")

    title = doc.add_heading("LAPORAN STATUS SENSOR", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Sistem Deteksi Kebakaran  —  Digenerate: {now_str}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("1. Informasi Umum", 1)
    _add_kv_table(doc, [
        ["Tanggal Laporan", datetime.now().strftime("%d %B %Y")],
        ["Waktu Generate", datetime.now().strftime("%H:%M:%S")],
        ["Kamera", camera_id],
        ["Dibuat oleh", "Sistem Deteksi Kebakaran — Chatbot K3"],
    ])

    cameras_to_report = (
        {camera_id: sensor_data[camera_id]}
        if camera_id != "all" and camera_id in sensor_data
        else sensor_data
    )

    doc.add_paragraph()
    doc.add_heading("2. Pembacaan Sensor Saat Ini", 1)

    for cam_id, d in cameras_to_report.items():
        doc.add_heading(f"Kamera: {cam_id}", 2)
        _add_data_table(doc, ["Sensor", "Nilai", "Satuan", "Status"], [
            ["MQ-2 (LPG/Asap)",         str(d.get("mq2", 0)),         "ppm", _gas_status(d.get("mq2", 0), 400)],
            ["MQ-4 (Metana)",            str(d.get("mq4", 0)),         "ppm", _gas_status(d.get("mq4", 0), 300)],
            ["MQ-5 (LPG/Gas Kota)",      str(d.get("mq5", 0)),         "ppm", _gas_status(d.get("mq5", 0), 300)],
            ["MQ-7 (CO)",                str(d.get("mq7", 0)),         "ppm", _gas_status(d.get("mq7", 0), 150)],
            ["MQ-135 (NH3/NOx/Benzena)", str(d.get("mq135", 0)),       "ppm", _gas_status(d.get("mq135", 0), 300)],
            ["Suhu",                     str(d.get("temperature", 0)), "°C",  _temp_status(d.get("temperature", 0))],
            ["Kelembapan",               str(d.get("humidity", 0)),    "% RH","Normal"],
        ])
        doc.add_paragraph()
        flame = d.get("flame_detected", False)
        lstm  = d.get("lstm_score")
        status_rows = [["Flame Sensor", "TERDETEKSI" if flame else "Tidak Terdeteksi"]]
        if lstm is not None:
            lstm_label = "ANOMALI TINGGI" if lstm >= 0.5 else "Normal"
            status_rows.append(["LSTM Anomaly Score", f"{lstm:.3f}  ({lstm_label})"])
        _add_kv_table(doc, status_rows)
        doc.add_paragraph()

    doc.add_heading("3. Kesimpulan", 1)
    all_data = list(cameras_to_report.values())
    if any(d.get("flame_detected") for d in all_data):
        doc.add_paragraph(
            "PERINGATAN: Flame sensor mendeteksi api! "
            "Segera lakukan prosedur evakuasi dan hubungi tim pemadam."
        )
    elif any((d.get("lstm_score") or 0) >= 0.5 for d in all_data):
        doc.add_paragraph(
            "PERHATIAN: Terdapat anomali tinggi pada LSTM score. "
            "Lakukan pemeriksaan fisik pada area yang terindikasi."
        )
    else:
        doc.add_paragraph("Kondisi sensor dalam batas normal. Tidak ada indikasi kebakaran aktif saat ini.")

    doc.add_paragraph()
    doc.add_heading("4. Rekomendasi K3", 1)
    for rec in _sensor_recs(cameras_to_report):
        doc.add_paragraph(rec, style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph("Laporan ini digenerate secara otomatis oleh Sistem Deteksi Kebakaran — PBL Politeknik Negeri Jakarta.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_incident_report(
    alerts: List[Dict],
    lstm_history: Dict,
    camera_id: str = "all",
) -> bytes:
    """Return DOCX bytes for an incident report."""
    doc = Document()
    now_str = datetime.now().strftime("%d %B %Y, %H:%M:%S")

    title = doc.add_heading("LAPORAN INSIDEN KEBAKARAN", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Sistem Deteksi Kebakaran  —  Digenerate: {now_str}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("1. Informasi Laporan", 1)
    _add_kv_table(doc, [
        ["Tanggal Laporan", datetime.now().strftime("%d %B %Y")],
        ["Waktu Generate", datetime.now().strftime("%H:%M:%S")],
        ["Kamera", camera_id],
        ["Jumlah Alert", str(len(alerts))],
        ["Dibuat oleh", "Sistem Deteksi Kebakaran — Chatbot K3"],
    ])

    doc.add_paragraph()
    doc.add_heading("2. Riwayat Alert / Insiden", 1)
    if alerts:
        rows = []
        for a in alerts:
            ts    = str(a.get("timestamp", a.get("time", "N/A")))
            cam   = str(a.get("camera_id", a.get("cam_id", "N/A")))
            level = str(a.get("level", a.get("severity", "N/A")))
            msg   = str(a.get("message", a.get("msg", str(a))))[:80]
            rows.append([ts, cam, level, msg])
        _add_data_table(doc, ["Waktu", "Kamera", "Level", "Pesan"], rows)
    else:
        doc.add_paragraph("Tidak ada alert yang tercatat dalam periode ini.")

    doc.add_paragraph()
    doc.add_heading("3. Riwayat Anomali LSTM", 1)
    history = lstm_history.get("history", [])
    if history:
        rows = []
        for h in history:
            ts    = str(h.get("timestamp", h.get("time", "N/A")))
            score = h.get("score", h.get("anomaly_score", 0)) or 0
            rows.append([ts, f"{float(score):.4f}", "ANOMALI" if float(score) >= 0.5 else "Normal"])
        _add_data_table(doc, ["Waktu", "Score", "Status"], rows)
    else:
        doc.add_paragraph("Tidak ada data riwayat LSTM untuk periode ini.")

    doc.add_paragraph()
    doc.add_heading("4. Analisis dan Kesimpulan", 1)
    high_alerts = [
        a for a in alerts
        if str(a.get("level", a.get("severity", ""))).upper()
        in {"HIGH", "CRITICAL", "TINGGI", "DARURAT"}
    ]
    anomaly_count = sum(
        1 for h in history
        if float(h.get("score", h.get("anomaly_score", 0)) or 0) >= 0.5
    )
    if high_alerts:
        doc.add_paragraph(f"Terdeteksi {len(high_alerts)} alert level tinggi. Diperlukan investigasi dan tindakan segera.")
    if anomaly_count > 0:
        doc.add_paragraph(f"Terdapat {anomaly_count} titik data anomali LSTM (score >= 0.5) dalam periode ini.")
    if not alerts and not anomaly_count:
        doc.add_paragraph("Tidak ditemukan insiden signifikan dalam periode yang dilaporkan.")

    doc.add_paragraph()
    doc.add_heading("5. Rekomendasi Tindak Lanjut", 1)
    for rec in _incident_recs(alerts, anomaly_count):
        doc.add_paragraph(rec, style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph("Laporan ini digenerate secara otomatis oleh Sistem Deteksi Kebakaran — PBL Politeknik Negeri Jakarta.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
